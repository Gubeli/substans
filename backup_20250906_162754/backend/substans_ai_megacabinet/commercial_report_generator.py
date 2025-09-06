


"""
Commercial Report Generator - Générateur de Rapports Commerciaux
Système pour la création automatisée de rapports commerciaux multi-formats (PDF, Word, Excel)
Intégré avec le système d'intelligence pour des rapports riches et pertinents
"""

import json
import time
from datetime import datetime
from typing import Dict, List, Any, Optional
from fpdf import FPDF
from openpyxl import Workbook
from docx import Document
import re

class CommercialReportGenerator:
    def __init__(self):
        self.name = "Commercial Report Generator"
        self.version = "1.0"
        self.report_templates = self._initialize_templates()
        
        print(f"🚀 {self.name} v{self.version} initialisé")
        print("✅ Templates de rapports chargés")
        print("✅ Export multi-format activé (PDF, Word, Excel)")

    def _initialize_templates(self) -> Dict[str, Any]:
        return {
            "market_analysis": {
                "name": "Analyse de Marché",
                "structure": [
                    {"section": "title_page", "title": "Analyse de Marché - {topic}"},
                    {"section": "executive_summary", "title": "Résumé Exécutif"},
                    {"section": "market_overview", "title": "Vue d'ensemble du Marché"},
                    {"section": "key_trends", "title": "Tendances Clés"},
                    {"section": "competitive_landscape", "title": "Paysage Concurrentiel"},
                    {"section": "swot_analysis", "title": "Analyse SWOT"},
                    {"section": "recommendations", "title": "Recommandations Stratégiques"},
                    {"section": "appendix", "title": "Annexe"}
                ],
                "formats": ["pdf", "docx"]
            },
            "competitive_intelligence": {
                "name": "Veille Concurrentielle",
                "structure": [
                    {"section": "title_page", "title": "Rapport de Veille Concurrentielle - {competitor}"},
                    {"section": "competitor_profile", "title": "Profil du Concurrent"},
                    {"section": "product_analysis", "title": "Analyse des Produits/Services"},
                    {"section": "market_positioning", "title": "Positionnement Marché"},
                    {"section": "recent_activities", "title": "Activités Récentes"},
                    {"section": "strategic_outlook", "title": "Perspectives Stratégiques"}
                ],
                "formats": ["pdf", "docx", "xlsx"]
            },
            "financial_report": {
                "name": "Rapport Financier",
                "structure": [
                    {"section": "title_page", "title": "Rapport Financier - {company}"},
                    {"section": "financial_summary", "title": "Résumé Financier"},
                    {"section": "income_statement", "title": "Compte de Résultat"},
                    {"section": "balance_sheet", "title": "Bilan"},
                    {"section": "cash_flow", "title": "Flux de Trésorerie"}
                ],
                "formats": ["pdf", "xlsx"]
            }
        }

    def generate_report(self, report_type: str, data: Dict[str, Any], output_format: str) -> Optional[str]:
        if report_type not in self.report_templates:
            print(f"Erreur: Type de rapport '{report_type}' non supporté.")
            return None

        template = self.report_templates[report_type]
        if output_format not in template["formats"]:
            print(f"Erreur: Format '{output_format}' non supporté pour ce type de rapport.")
            return None

        output_path = f"/home/ubuntu/substans_ai_megacabinet/reports/report_{report_type}_{int(time.time())}.{output_format}"

        if output_format == 'pdf':
            self._generate_pdf_report(template, data, output_path)
        elif output_format == 'docx':
            self._generate_docx_report(template, data, output_path)
        elif output_format == 'xlsx':
            self._generate_xlsx_report(template, data, output_path)
        
        return output_path

    def _generate_pdf_report(self, template: Dict[str, Any], data: Dict[str, Any], output_path: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for section in template["structure"]:
            title = section["title"].format(**data)
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt=title, ln=True, align='C')
            
            content_method = getattr(self, f"_get_{section['section']}_content", None)
            if content_method:
                content = content_method(data)
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, txt=content)
                pdf.ln(10)

        pdf.output(output_path)
        print(f"Rapport PDF généré: {output_path}")

    def _generate_docx_report(self, template: Dict[str, Any], data: Dict[str, Any], output_path: str):
        document = Document()

        for section in template["structure"]:
            title = section["title"].format(**data)
            document.add_heading(title, level=1)
            
            content_method = getattr(self, f"_get_{section['section']}_content", None)
            if content_method:
                content = content_method(data)
                document.add_paragraph(content)

        document.save(output_path)
        print(f"Rapport DOCX généré: {output_path}")

    def _generate_xlsx_report(self, template: Dict[str, Any], data: Dict[str, Any], output_path: str):
        workbook = Workbook()
        
        for i, section in enumerate(template["structure"]):
            title = section["title"].format(**data)
            sanitized_title = re.sub(r'[\\/*?:\[\]]', '', title)[:31]
            if i == 0:
                worksheet = workbook.active
                worksheet.title = sanitized_title
            else:
                worksheet = workbook.create_sheet(title=sanitized_title, index=i)


            content_method = getattr(self, f"_get_{section['section']}_data", None)
            if content_method:
                section_data = content_method(data)
                for row_index, row_data in enumerate(section_data, 1):
                    for col_index, cell_value in enumerate(row_data, 1):
                        worksheet.cell(row=row_index, column=col_index, value=str(cell_value))

        workbook.save(output_path)
        print(f"Rapport XLSX généré: {output_path}")

    # --- Méthodes de contenu pour les sections de rapport ---

    def _get_title_page_content(self, data: Dict[str, Any]) -> str:
        return f"Rapport généré par substans.ai\nDate: {datetime.now().strftime('%Y-%m-%d')}"

    def _get_executive_summary_content(self, data: Dict[str, Any]) -> str:
        return data.get("summary", "Résumé non disponible.")

    def _get_market_overview_content(self, data: Dict[str, Any]) -> str:
        return data.get("market_overview", "Vue d'ensemble non disponible.")

    def _get_key_trends_content(self, data: Dict[str, Any]) -> str:
        trends = data.get("trends", [])
        return "\n".join([f"- {trend}" for trend in trends])

    def _get_competitive_landscape_content(self, data: Dict[str, Any]) -> str:
        competitors = data.get("competitors", [])
        return "\n".join([f"- {c['name']}: {c['market_share']}%" for c in competitors])

    def _get_swot_analysis_content(self, data: Dict[str, Any]) -> str:
        swot = data.get("swot", {})
        return f"Forces: {', '.join(swot.get('strengths', []))}\nFaiblesses: {', '.join(swot.get('weaknesses', []))}\nOpportunités: {', '.join(swot.get('opportunities', []))}\nMenaces: {', '.join(swot.get('threats', []))}"

    def _get_recommendations_content(self, data: Dict[str, Any]) -> str:
        recommendations = data.get("recommendations", [])
        return "\n".join([f"- {rec}" for rec in recommendations])

    def _get_appendix_content(self, data: Dict[str, Any]) -> str:
        return data.get("appendix", "Pas d'annexes.")

    def _get_competitor_profile_data(self, data: Dict[str, Any]) -> List[List[Any]]:
        profile = data.get("profile", {})
        return [
            ["Nom", profile.get("name", "N/A")],
            ["Secteur", profile.get("sector", "N/A")],
            ["Chiffre d'affaires", profile.get("revenue", "N/A")]
        ]

# --- Test et Démonstration ---
if __name__ == '__main__':
    generator = CommercialReportGenerator()

    # Données de test
    market_data = {
        "topic": "IA Générative",
        "summary": "Le marché de l'IA générative est en pleine expansion.",
        "market_overview": "...",
        "trends": ["Adoption par les entreprises", "Modèles multimodaux"],
        "competitors": [{"name": "OpenAI", "market_share": 45}, {"name": "Google", "market_share": 30}],
        "swot": {"strengths": ["Innovation rapide"], "weaknesses": ["Coûts élevés"]},
        "recommendations": ["Investir dans la R&D"]
    }

    # Génération d'un rapport PDF
    generator.generate_report("market_analysis", market_data, "pdf")

    # Génération d'un rapport DOCX
    generator.generate_report("market_analysis", market_data, "docx")

    # Données de test pour la veille concurrentielle
    competitor_data = {
        "competitor": "OpenAI",
        "profile": {"name": "OpenAI", "sector": "IA", "revenue": "$2B"}
    }

    # Génération d'un rapport XLSX
    generator.generate_report("competitive_intelligence", competitor_data, "xlsx")




    def _get_financial_summary_content(self, data: Dict[str, Any]) -> str:
        summary = data.get("financial_summary", {})
        return f"Chiffre d'affaires: {summary.get('revenue', 'N/A')}\nBénéfice Net: {summary.get('net_profit', 'N/A')}\nMarge Brute: {summary.get('gross_margin', 'N/A')}%"

    def _get_income_statement_data(self, data: Dict[str, Any]) -> List[List[Any]]:
        statement = data.get("income_statement", {})
        return [
            ["Revenus", statement.get("revenue", 0)],
            ["Coût des marchandises vendues", statement.get("cogs", 0)],
            ["Bénéfice Brut", statement.get("gross_profit", 0)],
            ["Dépenses d'exploitation", statement.get("operating_expenses", 0)],
            ["Bénéfice d'exploitation", statement.get("operating_income", 0)],
            ["Impôts", statement.get("taxes", 0)],
            ["Bénéfice Net", statement.get("net_income", 0)]
        ]

    def _get_balance_sheet_data(self, data: Dict[str, Any]) -> List[List[Any]]:
        sheet = data.get("balance_sheet", {})
        return [
            ["Actifs Courants", sheet.get("current_assets", 0)],
            ["Actifs non courants", sheet.get("non_current_assets", 0)],
            ["Total Actifs", sheet.get("total_assets", 0)],
            ["Passifs courants", sheet.get("current_liabilities", 0)],
            ["Passifs non courants", sheet.get("non_current_liabilities", 0)],
            ["Total Passifs", sheet.get("total_liabilities", 0)],
            ["Capitaux Propres", sheet.get("equity", 0)]
        ]

    def _get_cash_flow_data(self, data: Dict[str, Any]) -> List[List[Any]]:
        flow = data.get("cash_flow", {})
        return [
            ["Flux de trésorerie d'exploitation", flow.get("operating", 0)],
            ["Flux de trésorerie d'investissement", flow.get("investing", 0)],
            ["Flux de trésorerie de financement", flow.get("financing", 0)],
            ["Variation nette de la trésorerie", flow.get("net_change", 0)]
        ]




    # Données de test pour le rapport financier
    financial_data = {
        "company": "TechCorp",
        "financial_summary": {"revenue": "$5B", "net_profit": "$1.2B", "gross_margin": 65},
        "income_statement": {"revenue": 5000, "cogs": 1750, "gross_profit": 3250, "operating_expenses": 1500, "operating_income": 1750, "taxes": 525, "net_income": 1225},
        "balance_sheet": {"current_assets": 2000, "non_current_assets": 8000, "total_assets": 10000, "current_liabilities": 1500, "non_current_liabilities": 2500, "total_liabilities": 4000, "equity": 6000},
        "cash_flow": {"operating": 2000, "investing": -1000, "financing": -500, "net_change": 500}
    }

    # Génération d'un rapport financier PDF
    generator.generate_report("financial_report", financial_data, "pdf")

    # Génération d'un rapport financier XLSX
    generator.generate_report("financial_report", financial_data, "xlsx")


