#!/usr/bin/env python3
"""
Advanced Report Generator - Système de Génération de Rapports Avancé
Correction et amélioration du système de génération de rapports
"""

import os
import json
import sqlite3
import datetime
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, asdict
from enum import Enum
import uuid
from pathlib import Path
import base64
from io import BytesIO

# Imports pour génération documents
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
except ImportError:
    print("⚠️ ReportLab non installé - Génération PDF limitée")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.chart import BarChart, Reference, LineChart, PieChart
except ImportError:
    print("⚠️ OpenPyXL non installé - Génération Excel limitée")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("⚠️ python-docx non installé - Génération Word limitée")

class ReportType(Enum):
    PERFORMANCE = "performance"
    SECURITY = "security"
    MISSIONS = "missions"
    INTELLIGENCE = "intelligence"
    FINANCIAL = "financial"
    QUALITY = "quality"
    TRENDS = "trends"
    CUSTOM = "custom"

class ReportFormat(Enum):
    PDF = "pdf"
    EXCEL = "xlsx"
    WORD = "docx"
    HTML = "html"
    JSON = "json"
    CSV = "csv"

class ReportFrequency(Enum):
    DAILY = "daily"
    WEEKLY = "weekly"
    MONTHLY = "monthly"
    QUARTERLY = "quarterly"
    YEARLY = "yearly"
    ON_DEMAND = "on_demand"

@dataclass
class ReportTemplate:
    """Template de rapport"""
    id: str
    name: str
    type: ReportType
    description: str
    sections: List[str]
    data_sources: List[str]
    default_format: ReportFormat
    parameters: Dict[str, Any]
    created_at: datetime.datetime

@dataclass
class ReportRequest:
    """Demande de génération de rapport"""
    id: str
    template_id: str
    name: str
    format: ReportFormat
    parameters: Dict[str, Any]
    requested_by: str
    requested_at: datetime.datetime
    status: str = "pending"
    file_path: Optional[str] = None
    generated_at: Optional[datetime.datetime] = None
    error_message: Optional[str] = None

@dataclass
class ReportSchedule:
    """Planification de rapport"""
    id: str
    template_id: str
    name: str
    frequency: ReportFrequency
    format: ReportFormat
    parameters: Dict[str, Any]
    next_execution: datetime.datetime
    last_execution: Optional[datetime.datetime] = None
    active: bool = True

class AdvancedReportGenerator:
    """Générateur de rapports avancé avec templates et export multi-format"""
    
    def __init__(self, base_path: str = "/home/ubuntu/substans_ai_megacabinet"):
        self.base_path = Path(base_path)
        self.reports_path = self.base_path / "reports"
        self.templates_path = self.base_path / "report_templates"
        self.db_path = self.base_path / "data" / "reports.db"
        
        # Créer les répertoires
        for path in [self.reports_path, self.templates_path, self.db_path.parent]:
            path.mkdir(parents=True, exist_ok=True)
        
        self._init_database()
        self._load_default_templates()
        
        # Données simulées pour les rapports
        self.sample_data = self._generate_sample_data()
    
    def _init_database(self):
        """Initialise la base de données des rapports"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS report_templates (
                    id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    type TEXT NOT NULL,
                    description TEXT,
                    sections TEXT,
                    data_sources TEXT,
                    default_format TEXT,
                    parameters TEXT,
                    created_at TIMESTAMP
                )
            """)
            
            conn.execute("""
                CREATE TABLE IF NOT EXISTS report_requests (
                    id TEXT PRIMARY KEY,
                    template_id TEXT NOT NULL,
                    name TEXT NOT NULL,
                    format TEXT NOT NULL,
                    parameters TEXT,
                    requested_by TEXT,
                    requested_at TIMESTAMP,
                    status TEXT DEFAULT 'pending',
                    file_path TEXT,
                    generated_at TIMESTAMP,
                    error_message TEXT
                )
            """)
            
            conn.execute("""
                CREATE TABLE IF NOT EXISTS report_schedules (
                    id TEXT PRIMARY KEY,
                    template_id TEXT NOT NULL,
                    name TEXT NOT NULL,
                    frequency TEXT NOT NULL,
                    format TEXT NOT NULL,
                    parameters TEXT,
                    next_execution TIMESTAMP,
                    last_execution TIMESTAMP,
                    active BOOLEAN DEFAULT TRUE
                )
            """)
            
            conn.commit()
    
    def _load_default_templates(self):
        """Charge les templates par défaut"""
        default_templates = [
            ReportTemplate(
                id="template_performance",
                name="Rapport de Performance",
                type=ReportType.PERFORMANCE,
                description="Rapport complet des performances système et agents",
                sections=["executive_summary", "system_metrics", "agent_performance", "trends", "recommendations"],
                data_sources=["system_monitor", "agent_metrics", "mission_stats"],
                default_format=ReportFormat.PDF,
                parameters={"period": "monthly", "include_charts": True},
                created_at=datetime.datetime.now()
            ),
            ReportTemplate(
                id="template_security",
                name="Rapport de Sécurité",
                type=ReportType.SECURITY,
                description="Analyse sécurité et conformité",
                sections=["security_overview", "threats", "compliance", "recommendations"],
                data_sources=["security_logs", "audit_trail", "compliance_checks"],
                default_format=ReportFormat.PDF,
                parameters={"include_incidents": True, "compliance_standards": ["GDPR", "SOX"]},
                created_at=datetime.datetime.now()
            ),
            ReportTemplate(
                id="template_missions",
                name="Rapport Missions",
                type=ReportType.MISSIONS,
                description="Analyse des missions et livrables",
                sections=["missions_overview", "success_rates", "deliverables", "client_satisfaction"],
                data_sources=["mission_data", "deliverable_stats", "client_feedback"],
                default_format=ReportFormat.EXCEL,
                parameters={"period": "quarterly", "include_details": True},
                created_at=datetime.datetime.now()
            ),
            ReportTemplate(
                id="template_intelligence",
                name="Rapport Intelligence",
                type=ReportType.INTELLIGENCE,
                description="Synthèse intelligence quotidienne et tendances",
                sections=["intelligence_summary", "key_insights", "trends", "recommendations"],
                data_sources=["daily_intelligence", "trend_analysis", "expert_insights"],
                default_format=ReportFormat.WORD,
                parameters={"sources_count": 290, "include_predictions": True},
                created_at=datetime.datetime.now()
            ),
            ReportTemplate(
                id="template_financial",
                name="Rapport Financier",
                type=ReportType.FINANCIAL,
                description="Analyse financière et ROI",
                sections=["financial_overview", "revenue_analysis", "cost_breakdown", "roi_metrics"],
                data_sources=["financial_data", "mission_revenues", "operational_costs"],
                default_format=ReportFormat.EXCEL,
                parameters={"currency": "EUR", "include_forecasts": True},
                created_at=datetime.datetime.now()
            )
        ]
        
        for template in default_templates:
            self._save_template(template)
    
    def _generate_sample_data(self) -> Dict[str, Any]:
        """Génère des données d'exemple pour les rapports"""
        return {
            "system_metrics": {
                "uptime": 99.95,
                "performance_score": 94.2,
                "security_score": 98.5,
                "compliance_score": 96.8,
                "active_agents": 32,
                "total_missions": 127,
                "success_rate": 98.5
            },
            "agent_performance": [
                {"name": "Senior Advisor", "performance": 98.2, "missions": 45, "satisfaction": 9.8},
                {"name": "Expert Finance & M&A", "performance": 96.5, "missions": 23, "satisfaction": 9.6},
                {"name": "Agent Proposition Commerciale", "performance": 94.8, "missions": 18, "satisfaction": 9.4},
                {"name": "Expert IA", "performance": 97.1, "missions": 31, "satisfaction": 9.7}
            ],
            "mission_stats": {
                "total": 127,
                "completed": 122,
                "active": 5,
                "success_rate": 98.5,
                "avg_duration": 12.5,
                "client_satisfaction": 9.2
            },
            "financial_data": {
                "revenue": 2500000,
                "costs": 1200000,
                "profit": 1300000,
                "roi": 108.3,
                "growth_rate": 45.2
            },
            "security_incidents": [
                {"date": "2025-09-01", "type": "Authentication", "severity": "Low", "resolved": True},
                {"date": "2025-09-02", "type": "Access", "severity": "Medium", "resolved": True}
            ],
            "intelligence_insights": [
                {"source": "Financial Times", "topic": "IA Enterprise", "relevance": 9.2},
                {"source": "MIT Technology Review", "topic": "HPC Trends", "relevance": 8.8},
                {"source": "McKinsey", "topic": "Digital Transformation", "relevance": 9.5}
            ]
        }
    
    def generate_report(self, template_id: str, format: ReportFormat, parameters: Dict[str, Any] = None) -> str:
        """Génère un rapport selon le template et format spécifiés"""
        template = self._get_template(template_id)
        if not template:
            raise ValueError(f"Template {template_id} non trouvé")
        
        # Créer la demande de rapport
        request_id = f"report_{uuid.uuid4().hex[:8]}"
        request = ReportRequest(
            id=request_id,
            template_id=template_id,
            name=f"{template.name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}",
            format=format,
            parameters=parameters or template.parameters,
            requested_by="system",
            requested_at=datetime.datetime.now()
        )
        
        try:
            # Générer le rapport selon le format
            if format == ReportFormat.PDF:
                file_path = self._generate_pdf_report(template, request)
            elif format == ReportFormat.EXCEL:
                file_path = self._generate_excel_report(template, request)
            elif format == ReportFormat.WORD:
                file_path = self._generate_word_report(template, request)
            elif format == ReportFormat.HTML:
                file_path = self._generate_html_report(template, request)
            elif format == ReportFormat.JSON:
                file_path = self._generate_json_report(template, request)
            else:
                raise ValueError(f"Format {format.value} non supporté")
            
            # Mettre à jour la demande
            request.status = "completed"
            request.file_path = file_path
            request.generated_at = datetime.datetime.now()
            
        except Exception as e:
            request.status = "error"
            request.error_message = str(e)
            file_path = None
        
        self._save_request(request)
        return file_path
    
    def _generate_pdf_report(self, template: ReportTemplate, request: ReportRequest) -> str:
        """Génère un rapport PDF"""
        filename = f"{request.name}.pdf"
        file_path = self.reports_path / filename
        
        # Créer le document PDF
        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Style personnalisé pour le titre
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        # Titre du rapport
        story.append(Paragraph(template.name, title_style))
        story.append(Spacer(1, 20))
        
        # Date de génération
        date_style = ParagraphStyle('DateStyle', parent=styles['Normal'], alignment=TA_CENTER)
        story.append(Paragraph(f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}", date_style))
        story.append(Spacer(1, 30))
        
        # Contenu selon le type de rapport
        if template.type == ReportType.PERFORMANCE:
            story.extend(self._generate_performance_content(styles))
        elif template.type == ReportType.SECURITY:
            story.extend(self._generate_security_content(styles))
        elif template.type == ReportType.MISSIONS:
            story.extend(self._generate_missions_content(styles))
        elif template.type == ReportType.INTELLIGENCE:
            story.extend(self._generate_intelligence_content(styles))
        elif template.type == ReportType.FINANCIAL:
            story.extend(self._generate_financial_content(styles))
        
        # Construire le PDF
        doc.build(story)
        return str(file_path)
    
    def _generate_performance_content(self, styles) -> List:
        """Génère le contenu du rapport de performance"""
        content = []
        data = self.sample_data
        
        # Résumé exécutif
        content.append(Paragraph("Résumé Exécutif", styles['Heading2']))
        content.append(Paragraph(
            f"La plateforme Substans.AI affiche d'excellentes performances avec un score global de "
            f"{data['system_metrics']['performance_score']}% et un uptime de {data['system_metrics']['uptime']}%. "
            f"Les {data['system_metrics']['active_agents']} agents ont traité {data['system_metrics']['total_missions']} "
            f"missions avec un taux de succès de {data['system_metrics']['success_rate']}%.",
            styles['Normal']
        ))
        content.append(Spacer(1, 20))
        
        # Métriques système
        content.append(Paragraph("Métriques Système", styles['Heading2']))
        metrics_data = [
            ['Métrique', 'Valeur', 'Statut'],
            ['Performance Globale', f"{data['system_metrics']['performance_score']}%", 'Excellent'],
            ['Sécurité', f"{data['system_metrics']['security_score']}%", 'Excellent'],
            ['Conformité', f"{data['system_metrics']['compliance_score']}%", 'Excellent'],
            ['Uptime', f"{data['system_metrics']['uptime']}%", 'Optimal']
        ]
        
        metrics_table = Table(metrics_data)
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        content.append(metrics_table)
        content.append(Spacer(1, 20))
        
        # Performance des agents
        content.append(Paragraph("Performance des Agents", styles['Heading2']))
        agent_data = [['Agent', 'Performance', 'Missions', 'Satisfaction']]
        for agent in data['agent_performance']:
            agent_data.append([
                agent['name'],
                f"{agent['performance']}%",
                str(agent['missions']),
                f"{agent['satisfaction']}/10"
            ])
        
        agent_table = Table(agent_data)
        agent_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        content.append(agent_table)
        
        return content
    
    def _generate_security_content(self, styles) -> List:
        """Génère le contenu du rapport de sécurité"""
        content = []
        data = self.sample_data
        
        content.append(Paragraph("Vue d'ensemble Sécurité", styles['Heading2']))
        content.append(Paragraph(
            f"Score de sécurité global : {data['system_metrics']['security_score']}%. "
            f"Conformité GDPR/SOX : {data['system_metrics']['compliance_score']}%. "
            f"{len(data['security_incidents'])} incidents traités ce mois.",
            styles['Normal']
        ))
        content.append(Spacer(1, 20))
        
        # Incidents de sécurité
        content.append(Paragraph("Incidents de Sécurité", styles['Heading2']))
        incident_data = [['Date', 'Type', 'Sévérité', 'Statut']]
        for incident in data['security_incidents']:
            incident_data.append([
                incident['date'],
                incident['type'],
                incident['severity'],
                'Résolu' if incident['resolved'] else 'En cours'
            ])
        
        incident_table = Table(incident_data)
        incident_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.red),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        content.append(incident_table)
        
        return content
    
    def _generate_missions_content(self, styles) -> List:
        """Génère le contenu du rapport missions"""
        content = []
        data = self.sample_data['mission_stats']
        
        content.append(Paragraph("Statistiques Missions", styles['Heading2']))
        content.append(Paragraph(
            f"Total missions : {data['total']} | Terminées : {data['completed']} | "
            f"Actives : {data['active']} | Taux succès : {data['success_rate']}%",
            styles['Normal']
        ))
        
        return content
    
    def _generate_intelligence_content(self, styles) -> List:
        """Génère le contenu du rapport intelligence"""
        content = []
        data = self.sample_data
        
        content.append(Paragraph("Synthèse Intelligence", styles['Heading2']))
        content.append(Paragraph(
            f"Analyse de {len(data['intelligence_insights'])} sources d'intelligence avec "
            f"relevance moyenne de 9.2/10.",
            styles['Normal']
        ))
        
        return content
    
    def _generate_financial_content(self, styles) -> List:
        """Génère le contenu du rapport financier"""
        content = []
        data = self.sample_data['financial_data']
        
        content.append(Paragraph("Analyse Financière", styles['Heading2']))
        content.append(Paragraph(
            f"Revenus : €{data['revenue']:,} | Coûts : €{data['costs']:,} | "
            f"Profit : €{data['profit']:,} | ROI : {data['roi']}%",
            styles['Normal']
        ))
        
        return content
    
    def _generate_excel_report(self, template: ReportTemplate, request: ReportRequest) -> str:
        """Génère un rapport Excel"""
        filename = f"{request.name}.xlsx"
        file_path = self.reports_path / filename
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Rapport"
        
        # Style pour les en-têtes
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Titre
        ws['A1'] = template.name
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')
        
        # Date
        ws['A2'] = f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        ws.merge_cells('A2:D2')
        
        # Données selon le type
        if template.type == ReportType.PERFORMANCE:
            self._add_performance_excel_data(ws, header_font, header_fill)
        elif template.type == ReportType.MISSIONS:
            self._add_missions_excel_data(ws, header_font, header_fill)
        elif template.type == ReportType.FINANCIAL:
            self._add_financial_excel_data(ws, header_font, header_fill)
        
        wb.save(str(file_path))
        return str(file_path)
    
    def _add_performance_excel_data(self, ws, header_font, header_fill):
        """Ajoute les données de performance à Excel"""
        data = self.sample_data
        
        # Métriques système
        ws['A4'] = "Métriques Système"
        ws['A4'].font = Font(size=14, bold=True)
        
        headers = ['Métrique', 'Valeur', 'Statut']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        metrics = [
            ['Performance', f"{data['system_metrics']['performance_score']}%", 'Excellent'],
            ['Sécurité', f"{data['system_metrics']['security_score']}%", 'Excellent'],
            ['Uptime', f"{data['system_metrics']['uptime']}%", 'Optimal']
        ]
        
        for row, metric in enumerate(metrics, 6):
            for col, value in enumerate(metric, 1):
                ws.cell(row=row, column=col, value=value)
    
    def _add_missions_excel_data(self, ws, header_font, header_fill):
        """Ajoute les données de missions à Excel"""
        data = self.sample_data['mission_stats']
        
        ws['A4'] = "Statistiques Missions"
        ws['A4'].font = Font(size=14, bold=True)
        
        stats = [
            ['Total Missions', data['total']],
            ['Missions Terminées', data['completed']],
            ['Missions Actives', data['active']],
            ['Taux de Succès', f"{data['success_rate']}%"],
            ['Durée Moyenne', f"{data['avg_duration']} jours"],
            ['Satisfaction Client', f"{data['client_satisfaction']}/10"]
        ]
        
        for row, (label, value) in enumerate(stats, 5):
            ws.cell(row=row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row, column=2, value=value)
    
    def _add_financial_excel_data(self, ws, header_font, header_fill):
        """Ajoute les données financières à Excel"""
        data = self.sample_data['financial_data']
        
        ws['A4'] = "Analyse Financière"
        ws['A4'].font = Font(size=14, bold=True)
        
        financials = [
            ['Revenus', f"€{data['revenue']:,}"],
            ['Coûts', f"€{data['costs']:,}"],
            ['Profit', f"€{data['profit']:,}"],
            ['ROI', f"{data['roi']}%"],
            ['Croissance', f"{data['growth_rate']}%"]
        ]
        
        for row, (label, value) in enumerate(financials, 5):
            ws.cell(row=row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row, column=2, value=value)
    
    def _generate_word_report(self, template: ReportTemplate, request: ReportRequest) -> str:
        """Génère un rapport Word"""
        filename = f"{request.name}.docx"
        file_path = self.reports_path / filename
        
        doc = Document()
        
        # Titre
        title = doc.add_heading(template.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenu selon le type
        if template.type == ReportType.INTELLIGENCE:
            self._add_intelligence_word_content(doc)
        elif template.type == ReportType.PERFORMANCE:
            self._add_performance_word_content(doc)
        
        doc.save(str(file_path))
        return str(file_path)
    
    def _add_intelligence_word_content(self, doc):
        """Ajoute le contenu intelligence au document Word"""
        data = self.sample_data
        
        doc.add_heading('Synthèse Intelligence Quotidienne', level=1)
        
        doc.add_paragraph(
            f"Cette synthèse présente l'analyse de {len(data['intelligence_insights'])} sources "
            f"d'intelligence avec une relevance moyenne de 9.2/10."
        )
        
        doc.add_heading('Insights Clés', level=2)
        for insight in data['intelligence_insights']:
            doc.add_paragraph(
                f"• {insight['source']}: {insight['topic']} (Relevance: {insight['relevance']}/10)",
                style='List Bullet'
            )
    
    def _add_performance_word_content(self, doc):
        """Ajoute le contenu performance au document Word"""
        data = self.sample_data
        
        doc.add_heading('Rapport de Performance', level=1)
        
        doc.add_paragraph(
            f"Score global de performance: {data['system_metrics']['performance_score']}% "
            f"avec {data['system_metrics']['active_agents']} agents actifs."
        )
    
    def _generate_html_report(self, template: ReportTemplate, request: ReportRequest) -> str:
        """Génère un rapport HTML"""
        filename = f"{request.name}.html"
        file_path = self.reports_path / filename
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{template.name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #366092; text-align: center; }}
                h2 {{ color: #4a90e2; border-bottom: 2px solid #4a90e2; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
                th {{ background-color: #366092; color: white; }}
                .metric {{ background-color: #f9f9f9; }}
            </style>
        </head>
        <body>
            <h1>{template.name}</h1>
            <p style="text-align: center;">Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
            
            <h2>Métriques Principales</h2>
            <table>
                <tr><th>Métrique</th><th>Valeur</th></tr>
                <tr class="metric"><td>Performance</td><td>{self.sample_data['system_metrics']['performance_score']}%</td></tr>
                <tr class="metric"><td>Sécurité</td><td>{self.sample_data['system_metrics']['security_score']}%</td></tr>
                <tr class="metric"><td>Uptime</td><td>{self.sample_data['system_metrics']['uptime']}%</td></tr>
            </table>
        </body>
        </html>
        """
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return str(file_path)
    
    def _generate_json_report(self, template: ReportTemplate, request: ReportRequest) -> str:
        """Génère un rapport JSON"""
        filename = f"{request.name}.json"
        file_path = self.reports_path / filename
        
        report_data = {
            "template": {
                "id": template.id,
                "name": template.name,
                "type": template.type.value
            },
            "generated_at": datetime.datetime.now().isoformat(),
            "data": self.sample_data
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)
        
        return str(file_path)
    
    def get_available_templates(self) -> List[ReportTemplate]:
        """Récupère tous les templates disponibles"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.execute("SELECT * FROM report_templates ORDER BY name")
            
            templates = []
            for row in cursor.fetchall():
                template = ReportTemplate(
                    id=row['id'],
                    name=row['name'],
                    type=ReportType(row['type']),
                    description=row['description'],
                    sections=json.loads(row['sections']),
                    data_sources=json.loads(row['data_sources']),
                    default_format=ReportFormat(row['default_format']),
                    parameters=json.loads(row['parameters']),
                    created_at=datetime.datetime.fromisoformat(row['created_at'])
                )
                templates.append(template)
            
            return templates
    
    def get_report_history(self, limit: int = 50) -> List[ReportRequest]:
        """Récupère l'historique des rapports générés"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.execute("""
                SELECT * FROM report_requests 
                ORDER BY requested_at DESC 
                LIMIT ?
            """, (limit,))
            
            requests = []
            for row in cursor.fetchall():
                request = ReportRequest(
                    id=row['id'],
                    template_id=row['template_id'],
                    name=row['name'],
                    format=ReportFormat(row['format']),
                    parameters=json.loads(row['parameters'] or '{}'),
                    requested_by=row['requested_by'],
                    requested_at=datetime.datetime.fromisoformat(row['requested_at']),
                    status=row['status'],
                    file_path=row['file_path'],
                    generated_at=datetime.datetime.fromisoformat(row['generated_at']) if row['generated_at'] else None,
                    error_message=row['error_message']
                )
                requests.append(request)
            
            return requests
    
    def _get_template(self, template_id: str) -> Optional[ReportTemplate]:
        """Récupère un template par son ID"""
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.execute("SELECT * FROM report_templates WHERE id = ?", (template_id,))
            row = cursor.fetchone()
            
            if not row:
                return None
            
            return ReportTemplate(
                id=row['id'],
                name=row['name'],
                type=ReportType(row['type']),
                description=row['description'],
                sections=json.loads(row['sections']),
                data_sources=json.loads(row['data_sources']),
                default_format=ReportFormat(row['default_format']),
                parameters=json.loads(row['parameters']),
                created_at=datetime.datetime.fromisoformat(row['created_at'])
            )
    
    def _save_template(self, template: ReportTemplate):
        """Sauvegarde un template"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                INSERT OR REPLACE INTO report_templates
                (id, name, type, description, sections, data_sources, default_format, parameters, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                template.id, template.name, template.type.value, template.description,
                json.dumps(template.sections), json.dumps(template.data_sources),
                template.default_format.value, json.dumps(template.parameters),
                template.created_at.isoformat()
            ))
    
    def _save_request(self, request: ReportRequest):
        """Sauvegarde une demande de rapport"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                INSERT OR REPLACE INTO report_requests
                (id, template_id, name, format, parameters, requested_by, requested_at, 
                 status, file_path, generated_at, error_message)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                request.id, request.template_id, request.name, request.format.value,
                json.dumps(request.parameters), request.requested_by,
                request.requested_at.isoformat(), request.status, request.file_path,
                request.generated_at.isoformat() if request.generated_at else None,
                request.error_message
            ))

# Instance globale
advanced_report_generator = AdvancedReportGenerator()

if __name__ == "__main__":
    # Test du générateur
    generator = AdvancedReportGenerator()
    
    # Générer un rapport de performance en PDF
    pdf_path = generator.generate_report("template_performance", ReportFormat.PDF)
    print(f"📊 Rapport PDF généré: {pdf_path}")
    
    # Générer un rapport missions en Excel
    excel_path = generator.generate_report("template_missions", ReportFormat.EXCEL)
    print(f"📈 Rapport Excel généré: {excel_path}")
    
    # Lister les templates
    templates = generator.get_available_templates()
    print(f"\n📋 Templates disponibles ({len(templates)}):")
    for template in templates:
        print(f"- {template.name} ({template.type.value}) - Format: {template.default_format.value}")

